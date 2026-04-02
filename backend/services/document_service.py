import csv
import os
import re


class DocumentService:
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc', '.rtf', '.md', '.csv'}
    TEXT_ENCODINGS = ('utf-8', 'utf-8-sig', 'gb18030', 'gbk')
    CONTENT_TYPE_MAP = {
        'text/plain': '.txt',
        'text/markdown': '.md',
        'text/csv': '.csv',
        'application/pdf': '.pdf',
        'application/msword': '.doc',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
        'application/rtf': '.rtf',
        'text/rtf': '.rtf',
    }

    @classmethod
    def infer_extension(cls, filename='', content_type=''):
        file_extension = os.path.splitext(filename or '')[1].lower()
        if file_extension in cls.SUPPORTED_EXTENSIONS:
            return file_extension
        return cls.CONTENT_TYPE_MAP.get((content_type or '').lower(), '')

    @classmethod
    def parse_document(cls, file_path, original_filename='', content_type=''):
        file_extension = cls.infer_extension(original_filename, content_type) or os.path.splitext(file_path)[1].lower()
        if file_extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {file_extension}")

        parser_map = {
            '.txt': cls.parse_txt,
            '.pdf': cls.parse_pdf,
            '.docx': cls.parse_docx,
            '.doc': cls.parse_doc,
            '.rtf': cls.parse_rtf,
            '.md': cls.parse_md,
            '.csv': cls.parse_csv,
        }

        raw_text = parser_map[file_extension](file_path)
        clean_text = cls.clean_markdown_text(raw_text) if file_extension == '.md' else cls.clean_text(raw_text)

        if not clean_text.strip():
            raise ValueError('文档解析完成，但未提取到有效文本内容')

        return clean_text

    @classmethod
    def read_text_file(cls, file_path):
        for encoding in cls.TEXT_ENCODINGS:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()

    @classmethod
    def clean_text(cls, text):
        if text is None:
            return ''

        text = str(text)
        text = text.replace('\ufeff', '').replace('\x00', '')
        text = text.replace('\r\n', '\n').replace('\r', '\n').replace('\xa0', ' ')
        text = re.sub(r'[ \t]+\n', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)
        text = re.sub(r'(?<=\w)-\n(?=\w)', '', text)

        lines = [line.strip() for line in text.split('\n')]
        merged_lines = []

        for line in lines:
            if not line:
                if merged_lines and merged_lines[-1] != '':
                    merged_lines.append('')
                continue

            is_list_like = bool(re.match(r'^(\d+[\.\)]|[-*•])\s+', line))
            if (
                merged_lines
                and merged_lines[-1] not in ('', None)
                and not re.search(r'[。！？；.!?;:]$', merged_lines[-1])
                and not is_list_like
            ):
                merged_lines[-1] = f"{merged_lines[-1]} {line}"
            else:
                merged_lines.append(line)

        return '\n'.join(merged_lines).strip()

    @classmethod
    def clean_markdown_text(cls, text):
        if text is None:
            return ''

        source = str(text).replace('\ufeff', '').replace('\x00', '')
        source = source.replace('\r\n', '\n').replace('\r', '\n').replace('\xa0', ' ')
        source = re.sub(r'^\s*---\n[\s\S]*?\n---\s*\n?', '', source, count=1)
        source = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', lambda m: f"图片：{m.group(1) or '未命名图片'} ({m.group(2)})", source)
        source = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', lambda m: f"{m.group(1)} ({m.group(2)})", source)

        code_blocks = []

        def stash_code_block(match):
            language = (match.group(1) or 'text').strip()
            code = match.group(2).strip()
            token = f"__CODE_BLOCK_{len(code_blocks)}__"
            block_text = f"\n代码块（{language}）：\n{code}\n"
            code_blocks.append((token, block_text))
            return token

        source = re.sub(r'```([\w+-]*)\n([\s\S]*?)```', stash_code_block, source)

        lines = []
        for raw_line in source.split('\n'):
            line = raw_line.rstrip()
            if not line.strip():
                if lines and lines[-1] != '':
                    lines.append('')
                continue

            if line.startswith('__CODE_BLOCK_') and line.endswith('__'):
                lines.append(line)
                continue

            line = re.sub(r'^\s{0,3}(#{1,6})\s*(.+)$', lambda m: f"标题：{m.group(2).strip()}", line)
            line = re.sub(r'^\s*>\s?', '引用：', line)
            line = re.sub(r'^\s*[-*+]\s+', '- ', line)
            line = re.sub(r'^\s*(\d+)\.\s+', r'\1. ', line)
            line = re.sub(r'\|', ' | ', line)
            line = re.sub(r'\s{2,}', ' ', line).strip()
            line = re.sub(r'^[#*_`>-]+\s*$', '', line).strip()

            if line:
                lines.append(line)
            elif lines and lines[-1] != '':
                lines.append('')

        normalized = '\n'.join(lines)
        normalized = re.sub(r'\n{3,}', '\n\n', normalized)

        for token, block_text in code_blocks:
            normalized = normalized.replace(token, block_text)

        return normalized.strip()

    @classmethod
    def parse_txt(cls, file_path):
        return cls.read_text_file(file_path)

    @classmethod
    def parse_md(cls, file_path):
        return cls.read_text_file(file_path)

    @classmethod
    def parse_csv(cls, file_path):
        lines = []
        for encoding in cls.TEXT_ENCODINGS:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as file:
                    reader = csv.reader(file)
                    for row in reader:
                        if any(cell.strip() for cell in row):
                            lines.append(' | '.join(cell.strip() for cell in row))
                break
            except UnicodeDecodeError:
                lines = []
                continue

        if not lines:
            with open(file_path, 'r', encoding='utf-8', errors='ignore', newline='') as file:
                reader = csv.reader(file)
                for row in reader:
                    if any(cell.strip() for cell in row):
                        lines.append(' | '.join(cell.strip() for cell in row))

        return '\n'.join(lines)

    @classmethod
    def parse_pdf(cls, file_path):
        try:
            import pdfplumber
        except ImportError as exc:
            raise ValueError('PDF 解析依赖 pdfplumber，请先安装该依赖') from exc

        segments = []
        with pdfplumber.open(file_path) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                page_text = (page.extract_text() or '').strip()
                if page_text:
                    segments.append(f'第 {page_index} 页\n{page_text}')

                for table_index, table in enumerate(page.extract_tables() or [], start=1):
                    table_rows = []
                    for row in table:
                        cells = [str(cell).strip() if cell is not None else '' for cell in row]
                        if any(cells):
                            table_rows.append(' | '.join(cells))
                    if table_rows:
                        segments.append(f'第 {page_index} 页表格 {table_index}\n' + '\n'.join(table_rows))

        return '\n\n'.join(segments)

    @classmethod
    def parse_docx(cls, file_path):
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError('DOCX 解析依赖 python-docx，请先安装该依赖') from exc

        doc = Document(file_path)
        segments = []

        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                segments.append(paragraph.text.strip())

        for table_index, table in enumerate(doc.tables, start=1):
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(' | '.join(cells))
            if table_rows:
                segments.append(f'表格 {table_index}\n' + '\n'.join(table_rows))

        return '\n'.join(segments)

    @classmethod
    def parse_doc(cls, file_path):
        try:
            return cls.parse_docx(file_path)
        except Exception as exc:
            raise ValueError('旧版 .doc 文档解析失败，建议转换为 .docx 后再上传') from exc

    @classmethod
    def parse_rtf(cls, file_path):
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError as exc:
            raise ValueError('RTF 解析依赖 striprtf，请先安装该依赖') from exc

        rtf_content = cls.read_text_file(file_path)
        return rtf_to_text(rtf_content)
